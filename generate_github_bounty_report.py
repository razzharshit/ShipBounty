from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, Polygon, Rectangle
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "report_assets"
OUT_DIR.mkdir(exist_ok=True)
REPORT_PATH = ROOT / "GitHub_Bounty_Dispenser_AI_Contribution_Score_Final_Report.docx"
WORKFLOW_PATH = OUT_DIR / "workflow_diagram.png"
GANTT_PATH = OUT_DIR / "gantt_chart.png"


BLUE = "#0B4F80"
TEAL = "#146B5C"
RED = "#B03535"
LIGHT = "#F4F7FB"
GREY = "#6B7280"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.line_spacing = 1.15
    styles["Normal"].paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Title", 22, "0B4F80"),
        ("Heading 1", 16, "0B4F80"),
        ("Heading 2", 13, "146B5C"),
        ("Heading 3", 11, "333333"),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold_start: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_start and text.startswith(bold_start):
        run = p.add_run(bold_start)
        run.bold = True
        p.add_run(text[len(bold_start) :])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, True)
        set_cell_shading(table.rows[0].cells[idx], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()


def generate_workflow_diagram() -> None:
    fig, ax = plt.subplots(figsize=(22, 5))
    ax.set_xlim(0, 24)
    ax.set_ylim(0, 6)
    ax.axis("off")

    def box(x, y, w, h, text, color=BLUE, txt="white"):
        ax.add_patch(Rectangle((x, y), w, h, facecolor=color, edgecolor="#222", linewidth=1.2))
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", color=txt, fontsize=8, wrap=True)

    def diamond(cx, cy, w, h, text):
        points = [(cx, cy + h / 2), (cx + w / 2, cy), (cx, cy - h / 2), (cx - w / 2, cy)]
        ax.add_patch(Polygon(points, closed=True, facecolor=TEAL, edgecolor="#222", linewidth=1.2))
        ax.text(cx, cy, text, ha="center", va="center", color="white", fontsize=8, wrap=True)

    def arrow(x1, y1, x2, y2, text=None):
        ax.add_patch(FancyArrowPatch((x1, y1), (x2, y2), arrowstyle="-|>", mutation_scale=10, lw=1.1, color="#111"))
        if text:
            ax.text((x1 + x2) / 2, (y1 + y2) / 2 + 0.18, text, fontsize=7, ha="center")

    box(0.2, 2.55, 1.3, 0.65, "GitHub PR\nOpened / Updated", "white", "#111")
    box(1.8, 2.55, 1.2, 0.65, "Webhook\nTriggered")
    diamond(4.0, 2.9, 1.6, 1.6, "Signature\nVerification\nHMAC SHA256")
    box(5.4, 2.55, 1.0, 0.65, "Fetch\nPR Data")
    box(6.8, 2.55, 1.35, 0.65, "Extract PR\nMetadata")
    box(8.5, 2.55, 1.55, 0.65, "GitHub App Auth\nJWT + Token")
    box(10.4, 2.55, 1.55, 0.65, "Fetch Files\nand Patch")
    box(12.3, 2.55, 1.55, 0.65, "Store PR and\nFile Changes")
    diamond(14.5, 2.9, 1.45, 1.35, "Duplicate\nEntry?")
    box(16.0, 3.55, 1.1, 0.55, "Ignore")
    box(17.45, 3.55, 1.35, 0.55, "Skip\nInsertion")

    ax.add_patch(Rectangle((16.0, 1.1), 5.2, 1.85, facecolor="#F8FAFC", edgecolor="#333", linewidth=1.0))
    ax.text(18.6, 2.78, "AI Evaluation and Fraud Check", ha="center", fontsize=8)
    box(16.35, 1.75, 1.45, 0.55, "AI Scoring\nEngine", "white", "#111")
    box(18.15, 2.25, 1.35, 0.45, "Quality\nScore")
    box(18.15, 1.7, 1.35, 0.45, "Relevance\nScore")
    box(18.15, 1.15, 1.35, 0.45, "Complexity\nScore")
    box(20.0, 1.7, 1.0, 0.55, "Total\nScore", "#E5E7EB", "#111")
    diamond(22.0, 1.95, 1.2, 1.2, "Fraud\nStatus")
    box(23.15, 2.65, 1.05, 0.5, "Reject", RED)
    box(23.15, 1.45, 1.05, 0.5, "Reward\nDeveloper")

    sequence = [
        (1.5, 2.875, 1.8, 2.875),
        (3.0, 2.875, 3.2, 2.875),
        (4.8, 2.875, 5.4, 2.875),
        (6.4, 2.875, 6.8, 2.875),
        (8.15, 2.875, 8.5, 2.875),
        (10.05, 2.875, 10.4, 2.875),
        (11.95, 2.875, 12.3, 2.875),
        (13.85, 2.875, 13.78, 2.875),
    ]
    for x1, y1, x2, y2 in sequence:
        arrow(x1, y1, x2, y2)
    arrow(14.95, 3.2, 16.0, 3.8, "Yes")
    arrow(17.1, 3.825, 17.45, 3.825)
    arrow(14.95, 2.55, 16.35, 2.0, "No")
    arrow(17.8, 2.0, 18.15, 2.48)
    arrow(17.8, 2.0, 18.15, 1.93)
    arrow(17.8, 2.0, 18.15, 1.38)
    arrow(19.5, 2.48, 20.0, 2.0)
    arrow(19.5, 1.93, 20.0, 2.0)
    arrow(19.5, 1.38, 20.0, 2.0)
    arrow(21.0, 1.98, 21.4, 1.98)
    arrow(22.6, 2.25, 23.15, 2.9, "Fraud")
    arrow(22.6, 1.75, 23.15, 1.7, "Genuine")
    arrow(4.0, 3.7, 23.15, 4.75, "Invalid")
    box(23.15, 4.45, 1.2, 0.55, "Reject\nRequest", RED)
    plt.tight_layout()
    fig.savefig(WORKFLOW_PATH, dpi=180, bbox_inches="tight")
    plt.close(fig)


def generate_gantt_chart() -> None:
    tasks = [
        ("Literature Review and Problem Analysis", 0, 3, "#AEC6CF"),
        ("System Architecture Design", 3, 2, "#FF6961"),
        ("GitHub API Integration", 5, 2, "#77DD77"),
        ("PR Data Processing and Feature Extraction", 7, 2, "#77DD77"),
        ("AI Model for PR Evaluation", 9, 4, "#77DD77"),
        ("Fraud Detection Module", 13, 2, "#CDB4DB"),
        ("Blockchain Smart Contract Development", 15, 3, "#FFB347"),
        ("Backend Development and API Integration", 16, 3, "#FFB347"),
        ("Frontend Dashboard Development", 18, 2, "#9ADBCF"),
        ("System Testing and Bug Fixes", 20, 2, "#9ADBCF"),
        ("Documentation and Final Report", 20, 2, "#9ADBCF"),
    ]
    fig, ax = plt.subplots(figsize=(14, 7))
    y_positions = list(range(len(tasks)))
    for y, (task, start, duration, color) in zip(y_positions, tasks):
        ax.barh(y, duration, left=start, height=0.52, color=color, edgecolor="white")
    ax.set_yticks(y_positions)
    ax.set_yticklabels([t[0] for t in tasks], fontsize=9)
    ax.invert_yaxis()
    ax.set_xlim(0, 22)
    ax.set_xticks(range(0, 23, 2))
    ax.set_xticklabels([f"Wk {i}" for i in range(0, 23, 2)])
    ax.grid(axis="x", linestyle="--", alpha=0.3)
    ax.set_xlabel("Project Timeline (22 Weeks)", fontweight="bold")
    ax.set_title("Work Plan - GitHub Bounty Dispenser with AI-based PR Evaluation and Blockchain Rewards", fontweight="bold")
    for spine in ["top", "right", "left"]:
        ax.spines[spine].set_visible(False)
    fig.tight_layout()
    fig.savefig(GANTT_PATH, dpi=180, bbox_inches="tight")
    plt.close(fig)


def add_title_page(doc: Document) -> None:
    for _ in range(2):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("A Major Project Report\n")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string("0B4F80")
    run = title.add_run("on\n")
    run.font.size = Pt(14)
    run = title.add_run("GitHub Bounty Dispenser with AI Contribution Score")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("146B5C")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\nSubmitted in partial fulfilment of the requirements for the award of degree\n").font.size = Pt(12)
    r = p.add_run("Bachelor of Technology / Bachelor of Engineering\n")
    r.bold = True
    r.font.size = Pt(13)
    p.add_run("in\nComputer Science and Engineering").font.size = Pt(12)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    rows = [
        ("Submitted By", "Name: __________________________\nRoll No: _______________________"),
        ("Submitted To", "Department of Computer Science and Engineering"),
        ("Guide / Supervisor", "__________________________"),
        ("Academic Year", "2025 - 2026"),
    ]
    for i, row in enumerate(rows):
        set_cell_text(table.rows[i].cells[0], row[0], True)
        set_cell_text(table.rows[i].cells[1], row[1])
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Institute / University Name\n").bold = True
    p.add_run("Department of Computer Science and Engineering")
    doc.add_page_break()


def add_front_matter(doc: Document) -> None:
    add_heading(doc, "Certificate", 1)
    add_para(
        doc,
        "This is to certify that the major project report entitled GitHub Bounty Dispenser with AI Contribution Score has been carried out by the student in partial fulfilment of the requirements for the award of the degree in Computer Science and Engineering. The work presented in this report is a record of original project development and study carried out under guidance during the academic session.",
    )
    doc.add_paragraph("\n\nGuide Signature: ____________________        Head of Department: ____________________")
    doc.add_page_break()

    add_heading(doc, "Declaration", 1)
    add_para(
        doc,
        "I hereby declare that the project report entitled GitHub Bounty Dispenser with AI Contribution Score is based on work carried out by me. The project has not been submitted elsewhere for the award of any degree or diploma. All external resources, tools, and references used during the project have been acknowledged appropriately.",
    )
    doc.add_paragraph("\n\nStudent Signature: ____________________        Date: ____________________")
    doc.add_page_break()

    add_heading(doc, "Acknowledgement", 1)
    add_para(
        doc,
        "I express my sincere gratitude to my project guide, faculty members, and department for their support and guidance throughout the development of this project. I am also thankful to the open-source community and the maintainers of FastAPI, PostgreSQL, SQLAlchemy, Alembic, GitHub Apps, and related tools that made the implementation of this system possible.",
    )
    add_para(
        doc,
        "This project helped me understand backend engineering, secure webhook processing, database persistence, API integration, artificial intelligence based contribution evaluation, and the future potential of blockchain-enabled reward systems.",
    )
    doc.add_page_break()

    add_heading(doc, "Abstract", 1)
    add_para(
        doc,
        "GitHub Bounty Dispenser with AI Contribution Score is a backend-oriented platform designed to automate the evaluation of GitHub pull requests and support objective bounty distribution. In traditional open-source workflows, maintainers manually inspect pull requests, estimate contribution value, identify low-effort changes, and decide rewards. This manual process becomes difficult when repositories receive many submissions or when financial incentives attract spam and duplicate work.",
    )
    add_para(
        doc,
        "The proposed system acts as an intelligent layer between GitHub repositories and future bounty payment infrastructure. It uses a GitHub App to receive pull request webhooks, verifies webhook signatures using HMAC SHA256, authenticates through GitHub App installation tokens, fetches pull request files and patches, and stores structured contribution data in PostgreSQL. The stored patch data becomes the foundation for an AI contribution scoring engine that will evaluate quality, relevance, complexity, and documentation value. The system also includes a planned fraud detection layer for duplicate work, suspicious velocity, and low-quality submissions.",
    )
    add_para(
        doc,
        "The current implementation completes the ingestion foundation: FastAPI backend, SQLAlchemy models, Alembic migrations, PostgreSQL storage, GitHub webhook processing, GitHub App authentication, pull request ingestion, pull request file ingestion, and patch storage. Future phases extend this foundation with AI scoring, fraud detection, blockchain smart contracts, bounty calculation, and a maintainer dashboard.",
    )
    doc.add_page_break()

    add_heading(doc, "Table of Contents", 1)
    contents = [
        "1. Introduction",
        "2. Problem Statement",
        "3. Objectives",
        "4. Literature Review",
        "5. Existing System",
        "6. Proposed System",
        "7. System Architecture",
        "8. Requirement Analysis",
        "9. Database Design",
        "10. Module Description",
        "11. Implementation Details",
        "12. AI Contribution Scoring Design",
        "13. Fraud Detection Design",
        "14. Blockchain Reward Distribution Plan",
        "15. Testing and Verification",
        "16. Results and Discussion",
        "17. Work Plan and Gantt Chart",
        "18. Conclusion and Future Scope",
        "19. References",
    ]
    for item in contents:
        doc.add_paragraph(item)
    doc.add_page_break()


def add_main_content(doc: Document) -> None:
    add_heading(doc, "1. Introduction", 1)
    add_para(
        doc,
        "Open-source software development depends on contributions from distributed developers. GitHub provides an efficient collaboration platform through issues, branches, pull requests, and reviews. However, rewarding contributors fairly remains a difficult problem. Maintainers usually decide contribution value manually, and this decision can be subjective, inconsistent, and time consuming.",
    )
    add_para(
        doc,
        "GitHub Bounty Dispenser with AI Contribution Score is designed to automate the first major step of this workflow: collecting and evaluating pull request data. The platform receives pull request events from GitHub, verifies them securely, stores pull request metadata, fetches changed files and patch content, and prepares the stored data for artificial intelligence based scoring.",
    )
    add_para(
        doc,
        "The project combines backend engineering, GitHub App integration, database design, webhook security, future AI scoring, fraud detection, and blockchain reward distribution into one long-term system.",
    )

    add_heading(doc, "2. Problem Statement", 1)
    add_para(
        doc,
        "Maintainers of open-source repositories face difficulty in objectively evaluating pull request quality and deciding bounty rewards. Manual evaluation is slow and subjective, while financial rewards can attract spam, duplicate patches, or low-effort submissions. A system is required that can automatically ingest pull request data, preserve code changes, calculate contribution scores, detect suspicious submissions, and support transparent reward distribution.",
    )

    add_heading(doc, "3. Objectives", 1)
    add_bullets(
        doc,
        [
            "To build a secure FastAPI backend that receives GitHub pull request webhooks.",
            "To verify webhook authenticity using HMAC SHA256 signature validation.",
            "To authenticate as a GitHub App and fetch pull request file changes.",
            "To store pull request metadata and patch data in PostgreSQL.",
            "To prepare a structured foundation for AI-based contribution scoring.",
            "To design fraud detection checks for duplicate, spam, and low-effort submissions.",
            "To plan future bounty calculation and blockchain-based reward distribution.",
        ],
    )

    add_heading(doc, "4. Literature Review", 1)
    add_para(
        doc,
        "Several bounty and open-source reward platforms exist, but most depend on manual issue assignment and subjective maintainer review. GitHub itself provides pull request and review workflows, but it does not provide a standardized contribution quality score. Existing continuous integration tools validate builds and tests, but they do not measure broader contribution value such as relevance, complexity, maintainability, and documentation quality.",
    )
    add_para(
        doc,
        "Recent progress in large language models makes it possible to analyze source code patches, summarize intent, detect boilerplate patterns, and compare changes against issue requirements. Combining this capability with secure GitHub App ingestion can create a more objective layer for contribution evaluation.",
    )

    add_heading(doc, "5. Existing System", 1)
    add_bullets(
        doc,
        [
            "Maintainers manually review pull requests and estimate contribution value.",
            "Reward allocation is often subjective and may vary between reviewers.",
            "Duplicate or low-effort pull requests can waste reviewer time.",
            "Patch data is not always stored in a structured format for later scoring.",
            "Fraud detection and bounty payment workflows are usually separate systems.",
        ],
    )

    add_heading(doc, "6. Proposed System", 1)
    add_para(
        doc,
        "The proposed system introduces an automated backend that receives GitHub pull request events, verifies security signatures, fetches pull request metadata and files, stores data in a normalized database, and prepares the information for AI scoring. The AI scoring layer will evaluate quality, relevance, complexity, and documentation value. A fraud detection layer will flag duplicate work and suspicious patterns. In future phases, valid scores will be passed to a bounty calculator and blockchain smart contract for transparent reward distribution.",
    )
    add_table(
        doc,
        ["Component", "Purpose"],
        [
            ["GitHub App", "Receives repository-level permissions and webhook events."],
            ["FastAPI Webhook Receiver", "Validates and processes pull request events."],
            ["GitHub API Client", "Fetches PR files, additions, deletions, and patches."],
            ["PostgreSQL Database", "Stores users, repositories, pull requests, files, and scores."],
            ["AI Scoring Engine", "Evaluates contribution quality and produces score components."],
            ["Fraud Detection Layer", "Flags duplicate, spam, or risky submissions."],
            ["Bounty Distribution Layer", "Calculates and dispenses rewards in future phases."],
        ],
    )

    add_heading(doc, "7. System Architecture", 1)
    add_para(
        doc,
        "The high-level architecture begins at a GitHub repository. When a contributor opens or updates a pull request, GitHub sends a webhook event to the FastAPI backend. The backend verifies the payload signature, extracts the event action, stores pull request data, generates a GitHub App installation token, fetches file-level changes, and stores patch data. The stored data is later used by the AI evaluation layer and fraud detection layer.",
    )
    doc.add_picture(str(WORKFLOW_PATH), width=Inches(7.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("Figure 1: Workflow of GitHub Bounty Dispenser with AI Contribution Score")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "8. Requirement Analysis", 1)
    add_heading(doc, "8.1 Functional Requirements", 2)
    add_bullets(
        doc,
        [
            "The system shall receive GitHub webhook events at POST /webhook/github.",
            "The system shall verify webhook signatures before processing payload data.",
            "The system shall process only pull_request opened and synchronize actions.",
            "The system shall ignore unrelated events safely without crashing.",
            "The system shall store pull request metadata in the database.",
            "The system shall fetch changed files and patch data using the GitHub API.",
            "The system shall store file-level changes in a pull_request_files table.",
            "The system shall prevent duplicate file insertion for repeated webhooks.",
        ],
    )
    add_heading(doc, "8.2 Non-Functional Requirements", 2)
    add_bullets(
        doc,
        [
            "Security: payload authenticity must be verified using HMAC SHA256.",
            "Reliability: irrelevant or malformed events must not break the server.",
            "Scalability: service and model layers should support future scoring modules.",
            "Maintainability: business logic should remain separated from API routes.",
            "Observability: webhook flow and file ingestion should be clearly logged.",
        ],
    )

    add_heading(doc, "9. Database Design", 1)
    add_para(
        doc,
        "The database uses PostgreSQL with SQLAlchemy ORM models and Alembic migrations. It stores contributor identities, repositories, pull request metadata, file-level patches, and future scoring results.",
    )
    add_table(
        doc,
        ["Table", "Purpose", "Important Fields"],
        [
            ["users", "Stores contributor details.", "id, github_id, username"],
            ["repositories", "Stores tracked repositories.", "id, github_repo_id, name"],
            ["pull_requests", "Stores PR metadata.", "id, github_pr_id, title, state, additions, deletions, changed_files"],
            ["pull_request_files", "Stores file-level changes.", "id, pr_id, filename, additions, deletions, patch"],
            ["scores", "Stores future AI evaluation output.", "quality_score, relevance_score, complexity_score, total_score"],
        ],
    )

    add_heading(doc, "10. Module Description", 1)
    add_table(
        doc,
        ["Module", "Responsibility"],
        [
            ["api", "Contains FastAPI routes and request-response handling."],
            ["services", "Contains business logic such as PR creation and file persistence."],
            ["models", "Defines SQLAlchemy database entities."],
            ["schemas", "Defines Pydantic validation and serialization models."],
            ["db", "Configures database engine and session factory."],
            ["github/webhook.py", "Receives, validates, filters, and processes GitHub webhook events."],
            ["github/auth.py", "Generates GitHub App JWTs and installation tokens."],
            ["github/client.py", "Calls GitHub API endpoints such as pull request files."],
        ],
    )

    add_heading(doc, "11. Implementation Details", 1)
    add_heading(doc, "11.1 Webhook Processing", 2)
    add_numbered(
        doc,
        [
            "GitHub sends a pull_request event to the FastAPI endpoint.",
            "The backend reads the raw request body.",
            "The HMAC SHA256 signature is compared with the expected signature.",
            "The event type and action are printed for debugging.",
            "Only opened and synchronize pull request actions are processed.",
            "Pull request metadata is extracted and stored.",
            "GitHub App authentication is performed using installation_id.",
            "Pull request files and patches are fetched and persisted.",
        ],
    )
    add_heading(doc, "11.2 File Ingestion", 2)
    add_para(
        doc,
        "File ingestion is a critical part of the system because future AI evaluation depends on patch data. The GitHub API endpoint GET /repos/{owner}/{repo}/pulls/{number}/files returns filename, additions, deletions, and patch data. The backend transforms this response into structured file records and stores them through the service layer.",
    )
    add_para(
        doc,
        "Duplicate prevention is handled using a unique constraint on pr_id and filename. If the same webhook is delivered again, already stored files are skipped and the stored count becomes zero. This makes the ingestion process idempotent.",
    )

    add_heading(doc, "12. AI Contribution Scoring Design", 1)
    add_para(
        doc,
        "The AI contribution scoring engine is planned as Phase 2. It will read stored patches, pull request metadata, issue descriptions, and repository context. The scoring engine will produce multiple score components rather than a single opaque number, making the evaluation more explainable for maintainers and contributors.",
    )
    add_table(
        doc,
        ["Score Category", "Evaluation Criteria", "Suggested Weight"],
        [
            ["Quality Score", "Correctness, readability, maintainability, test impact.", "35%"],
            ["Relevance Score", "How well the PR solves the issue or requirement.", "25%"],
            ["Complexity Score", "Effort level, algorithmic difficulty, affected files.", "20%"],
            ["Documentation Score", "PR description, comments, clarity, usage notes.", "10%"],
            ["Reliability Score", "Build safety, edge cases, regression risk.", "10%"],
        ],
    )
    add_para(
        doc,
        "The total score can be calculated as a weighted aggregation of the score components. The score should also include a short explanation generated by the AI model so that maintainers can understand why a contribution received a particular rating.",
    )

    add_heading(doc, "13. Fraud Detection Design", 1)
    add_bullets(
        doc,
        [
            "Duplicate Work Detection: compare patch similarity between submissions.",
            "Spam Detection: identify tiny changes, repeated boilerplate, and unrelated modifications.",
            "Velocity Detection: flag accounts submitting too many PRs in a short period.",
            "Account Risk Analysis: consider account age and repository history.",
            "AI-generated Boilerplate Detection: detect suspicious repetitive patterns.",
        ],
    )
    add_para(
        doc,
        "Fraud detection does not automatically reject all flagged pull requests. Instead, suspicious submissions can be marked for review. This protects maintainers from spam while still allowing human judgement where required.",
    )

    add_heading(doc, "14. Blockchain Reward Distribution Plan", 1)
    add_para(
        doc,
        "The long-term bounty distribution layer will connect contribution scores to transparent payments. After a pull request is evaluated and marked genuine, a bounty calculator can determine the reward amount. A smart contract on Polygon or another blockchain network can then distribute rewards to contributor wallet addresses.",
    )
    add_table(
        doc,
        ["Step", "Description"],
        [
            ["Score Generation", "AI engine produces quality, relevance, complexity, and total scores."],
            ["Fraud Check", "Suspicious or duplicate work is flagged before payment."],
            ["Bounty Calculation", "Reward is calculated based on score and bounty pool."],
            ["Smart Contract Execution", "Contract transfers reward to developer wallet."],
            ["Audit Trail", "Transaction history provides transparent payout proof."],
        ],
    )

    add_heading(doc, "15. Testing and Verification", 1)
    add_table(
        doc,
        ["Test Case", "Expected Result", "Status"],
        [
            ["Valid pull_request.opened webhook", "PR metadata is stored and files are fetched.", "Implemented"],
            ["Valid pull_request.synchronize webhook", "Existing PR is reused and new file data is handled.", "Implemented"],
            ["Invalid webhook signature", "Request returns 401 Unauthorized.", "Implemented"],
            ["installation or ping event", "Event is ignored safely.", "Implemented"],
            ["pull_request.closed action", "Action is ignored safely.", "Implemented"],
            ["Repeated webhook delivery", "Duplicate file insertion is skipped.", "Implemented"],
            ["Missing GitHub file-fetch fields", "Webhook returns safely and logs missing data.", "Implemented"],
        ],
    )
    add_para(
        doc,
        "The current system has been designed to return safely for irrelevant events and file ingestion errors. Debug logs show event name, action, pull request processing, file count fetched, and file count stored.",
    )

    add_heading(doc, "16. Results and Discussion", 1)
    add_para(
        doc,
        "Phase 1 successfully establishes the data ingestion foundation required for AI-based contribution evaluation. The system can receive GitHub pull request events, verify signatures, persist pull request metadata, authenticate with GitHub App credentials, fetch pull request files, and store patch data. The stored patch is the most important input for future code-quality analysis.",
    )
    add_para(
        doc,
        "The result is a working backend foundation that can be extended without redesigning the core ingestion pipeline. Future scoring, fraud detection, and bounty distribution modules can build directly on the pull_requests and pull_request_files tables.",
    )

    add_heading(doc, "17. Work Plan and Gantt Chart", 1)
    add_para(
        doc,
        "The project plan is divided into six major phases: research, system design, AI model development, fraud detection, blockchain and backend integration, and frontend/testing/documentation. The complete plan spans 22 weeks.",
    )
    doc.add_picture(str(GANTT_PATH), width=Inches(7.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("Figure 2: 22-week Gantt chart for project execution")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "18. Conclusion and Future Scope", 1)
    add_para(
        doc,
        "GitHub Bounty Dispenser with AI Contribution Score provides a practical foundation for an automated contribution marketplace. The current implementation solves the first important challenge: securely collecting and storing pull request data and patch information. This enables future AI evaluation, fraud detection, and reward distribution.",
    )
    add_para(
        doc,
        "Future work includes implementing the scoring engine, connecting issue context to PR evaluation, building a maintainer dashboard, adding fraud analytics, deploying smart contracts, integrating wallet addresses, and enabling transparent bounty payouts.",
    )

    add_heading(doc, "19. References", 1)
    refs = [
        "FastAPI Documentation: https://fastapi.tiangolo.com/",
        "GitHub Apps Documentation: https://docs.github.com/en/apps",
        "GitHub REST API Documentation: https://docs.github.com/en/rest",
        "SQLAlchemy Documentation: https://docs.sqlalchemy.org/",
        "Alembic Documentation: https://alembic.sqlalchemy.org/",
        "PostgreSQL Documentation: https://www.postgresql.org/docs/",
        "Python requests Documentation: https://requests.readthedocs.io/",
        "Solidity Documentation: https://docs.soliditylang.org/",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Number")

    add_heading(doc, "Appendix A: Key Webhook Logs", 1)
    p = doc.add_paragraph()
    p.add_run(
        "EVENT: pull_request\nACTION: synchronize\nPR processing started\nProcessing PR: <title> by <author>\nFiles fetched: 1\nStored files: 1"
    ).font.name = "Courier New"


def build_report() -> None:
    generate_workflow_diagram()
    generate_gantt_chart()

    doc = Document()
    style_document(doc)
    add_title_page(doc)
    add_front_matter(doc)
    add_main_content(doc)

    for section in doc.sections:
        footer = section.footer
        if footer.paragraphs:
            add_page_number(footer.paragraphs[0])

    doc.save(REPORT_PATH)
    print(REPORT_PATH)


if __name__ == "__main__":
    build_report()
