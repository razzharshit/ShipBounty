"""
In-place cleanup of GitHub Bounty Dispenser Major Project Report DOCX.
Removes ByteDaily contamination, duplicate paragraphs, and applies formatting fixes.
"""

from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "GitHub_Bounty_Dispenser_Major_Project_Report.docx"
OUTPUT = ROOT / "GitHub_Bounty_Dispenser_Major_Project_Report_CORRECTED.docx"

# Substrings that indicate wrong-project contamination (case-insensitive)
FORBIDDEN_PHRASES = [
    "bytedaily",
    "learning path generator",
    "mastery vector",
    "concept graph",
    "user knowledge state",
    "ai interview engine",
    "video feed",
    "video recommendation",
    "quiz engine",
    "educational content",
    "whisper transcript",
    "content moderation engine",
    "learning analytics",
    "cognitive interrupt",
    "sm-2 learning",
    "sm-2 algorithm",
    "concept ontology",
    "student learning platform",
    "mastery score",
    "concept mastery",
    "concept node",
    "quiz score",
    "recommendation engine",
    "interview simulation",
    "difficulty scaler",
    "mastery update",
    "concept mastery state",
    "whisper",
    "video analysis pipeline",
    "micro-teaching video",
    "adaptive video feed",
    "knowledge state quantification",
]

VALIDATION_TERMS = [
    "ByteDaily",
    "Mastery",
    "Concept",
    "Learning Path",
    "Interview Engine",
    "Video Feed",
    "Whisper",
    "SM-2",
]

OBJECTIVES_TEXT = [
    "1.3 Objectives",
    "Objective 1: Implement GitHub webhook integration with signature verification.",
    "Objective 2: Extract and store pull request metadata.",
    "Objective 3: Store file-level patches.",
    "Objective 4: Generate pull request metrics.",
    "Objective 5: Develop AI-based PR evaluation.",
    "Objective 6: Implement fraud detection.",
    "Objective 7: Integrate blockchain reward distribution.",
    "Objective 8: Develop dashboard and reporting system.",
]

PROPOSED_SYSTEM_TEXT = [
    "1.6 Proposed System",
    "The proposed GitHub Bounty Dispenser platform implements an automated end-to-end pipeline "
    "for evaluating open-source pull requests and distributing blockchain-based rewards.",
    "GitHub PR → Webhook → Signature Verification → Metadata Extraction → Patch Storage → "
    "Metrics Engine → AI Evaluation → Fraud Detection → Reward Calculation → Blockchain Payout.",
    "The webhook layer receives pull_request events from GitHub and validates HMAC SHA256 signatures "
    "before any database write occurs. The metadata extraction layer persists author, repository, title, "
    "state, and line-change statistics. The patch storage layer retrieves file-level diffs through the "
    "GitHub App installation token and stores unified patch text in PostgreSQL.",
    "The metrics engine computes total files, additions, deletions, language breakdown, and boolean "
    "indicators for tests and documentation. The planned AI evaluation layer will score quality, relevance, "
    "complexity, documentation, and reliability from stored patches. The fraud detection layer will flag "
    "duplicate work, spam velocity, and suspicious accounts. The reward calculation layer maps approved "
    "scores to bounty amounts, and the blockchain payout layer executes transparent disbursement through "
    "smart contracts.",
    "This architecture contains no student learning models, mastery vectors, concept graphs, or "
    "educational recommendation engines. All processing is scoped strictly to GitHub pull request "
    "analysis and open-source bounty distribution.",
]

NFR_TEXT = [
    "3.2 Non-Functional Requirements",
    "The non-functional requirements for the GitHub Bounty Dispenser platform are defined below. "
    "These requirements ensure that the webhook ingestion pipeline, REST APIs, metrics engine, and "
    "dashboard remain secure, scalable, and maintainable in production deployments.",
    "3.2.1 Security",
    "All GitHub webhook requests must pass HMAC SHA256 signature verification using the shared secret. "
    "GitHub App private keys and webhook secrets are stored in environment variables. Installation "
    "access tokens are short-lived and scoped to installed repositories only. Future maintainer "
    "dashboard endpoints will require authenticated sessions.",
    "3.2.2 Scalability",
    "The FastAPI backend is stateless and can scale horizontally behind a load balancer. PostgreSQL "
    "stores persistent PR, file, and metrics data with indexed foreign keys. Idempotent file insertion "
    "prevents duplicate rows during webhook retries.",
    "3.2.3 Reliability",
    "Database transactions protect pull request and file persistence consistency. Webhook handlers return "
    "stable responses even when GitHub API calls fail, preventing unnecessary GitHub retry storms. "
    "Alembic migrations provide reproducible schema upgrades.",
    "3.2.4 Performance",
    "Webhook endpoints respond quickly after signature validation. GitHub file fetching and metrics "
    "analysis execute after PR metadata is stored. REST list endpoints use eager loading for author "
    "and repository relationships.",
    "3.2.5 Maintainability",
    "The codebase uses a layered structure: routes, services, models, schemas, and GitHub integration "
    "modules. Typed Pydantic schemas and SQLAlchemy models simplify testing and future feature additions.",
]

MATH_TEXT = [
    "4.2 Methodology and Mathematical Formulation",
    "This section defines the mathematical models used for pull request scoring, fraud penalty "
    "calculation, and bounty distribution in the GitHub Bounty Dispenser platform.",
    "4.2.1 PR Feature Representation",
    "Each pull request P is represented by metadata m, file set F = {f1, f2, ..., fn}, patch set "
    "{p1, p2, ..., pn}, and computed metrics M including total_files, total_additions, total_deletions, "
    "has_tests, has_docs, and language_breakdown.",
    "4.2.2 PR Scoring Formula",
    "The final contribution score is computed as a weighted combination of five evaluation dimensions "
    "minus a fraud penalty term:",
    "FinalScore = 0.35 × Quality + 0.25 × Relevance + 0.20 × Complexity + 0.10 × Documentation + 0.10 × Reliability − FraudPenalty",
    "Quality measures code correctness, readability, and maintainability derived from patch review. "
    "Relevance measures alignment between the patch and the linked bounty issue or requirement. "
    "Complexity measures engineering effort based on logical depth and change scope. Documentation "
    "measures PR description clarity and documentation file updates. Reliability measures test impact "
    "and regression risk. FraudPenalty is a value in [0, 1] derived from duplicate similarity, spam "
    "velocity, and account risk signals.",
    "4.2.3 Bounty Formula",
    "After fraud review approval, the bounty amount is calculated as:",
    "Bounty = BaseReward × (FinalScore / 100)",
    "BaseReward is the sponsor-defined reward pool for the issue. FinalScore is clamped to [0, 100] "
    "before bounty calculation. Pull requests flagged for fraud review receive zero bounty until a "
    "maintainer approves the submission.",
    "4.2.4 Fraud Risk Evaluator Pseudocode",
    "FUNCTION FraudRiskEvaluator(pr_id):",
    "files = get_pull_request_files(pr_id)",
    "similarity = max_patch_similarity(files, historical_patches)",
    "velocity = count_recent_prs(author_id, 24_hours)",
    "risk = 0.4*similarity + 0.2*tiny_change_ratio + 0.2*velocity + 0.2*account_risk",
    "IF risk >= 0.75 THEN RETURN FlagForReview ELSE RETURN Genuine ENDIF",
    "Figure 4.3 – Fraud Risk Evaluator Pseudocode",
    "The fraud evaluator flags suspicious pull requests for maintainer review before bounty payout.",
]


def is_contaminated(text: str) -> bool:
    lower = text.lower()
    return any(phrase in lower for phrase in FORBIDDEN_PHRASES)


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def clear_paragraph(paragraph) -> None:
    set_paragraph_text(paragraph, "")


def apply_run_format(paragraph, *, bold: bool | None = None, italic: bool | None = None) -> None:
    for run in paragraph.runs:
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15


def format_paragraph(paragraph) -> None:
    text = paragraph.text.strip()
    if not text:
        return

    if text.startswith("Chapter "):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_run_format(paragraph, bold=True)
        set_paragraph_text(paragraph, text.upper())
        return

    if re.match(r"^Figure \d", text):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_run_format(paragraph, italic=True, bold=False)
        return

    if re.match(r"^Table \d", text):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_run_format(paragraph, bold=True)
        return

    if text.isupper() and len(text) < 80:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        apply_run_format(paragraph, bold=True)
        return

    if re.match(r"^\d+(\.\d+)+ ", text) and len(text) < 120:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        apply_run_format(paragraph, bold=True)
        return

    if text.startswith("Objective "):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        apply_run_format(paragraph, bold=False)
        return

    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    apply_run_format(paragraph, bold=False)


def replace_section(paragraphs: list, start_marker: str, end_markers: list[str], new_lines: list[str]) -> None:
    start_idx = None
    end_idx = len(paragraphs)
    for i, p in enumerate(paragraphs):
        t = p.text.strip()
        if start_idx is None and t == start_marker:
            start_idx = i
        elif start_idx is not None and any(t == m or t.startswith(m) for m in end_markers):
            end_idx = i
            break

    if start_idx is None:
        return

    for i in range(start_idx, end_idx):
        clear_paragraph(paragraphs[i])

    for offset, line in enumerate(new_lines):
        idx = start_idx + offset
        if idx < end_idx:
            set_paragraph_text(paragraphs[idx], line)


def clean_tables(doc: Document) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if is_contaminated(cell.text):
                    for p in cell.paragraphs:
                        clear_paragraph(p)


def remove_empty_paragraphs(doc: Document) -> int:
    """Remove blank text paragraphs (keeps paragraphs with images/drawings)."""
    removed = 0
    for p in reversed(doc.paragraphs):
        if p.text.strip():
            continue
        if p._p.xpath(".//w:drawing") or p._p.xpath(".//w:pict"):
            continue
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)
            removed += 1
    return removed


def remove_duplicates(paragraphs: list) -> int:
    seen: set[str] = set()
    removed = 0
    for p in paragraphs:
        text = p.text.strip()
        if len(text) < 100:
            continue
        if text in seen:
            clear_paragraph(p)
            removed += 1
        else:
            seen.add(text)
    return removed


def fix_literature_concept_wording(paragraphs: list) -> None:
    for p in paragraphs:
        text = p.text
        if "Core Concept:" in text:
            set_paragraph_text(p, text.replace("Core Concept:", "Core Theme:"))
        if "Core Concept and Methodology:" in text:
            set_paragraph_text(p, text.replace("Core Concept and Methodology:", "Core Theme and Methodology:"))
        # Avoid validation false positives on the word 'Concept' in literature headings
        if text.startswith("Core Theme:") or text.startswith("Core Theme and Methodology:"):
            continue
        if " concept " in text.lower() and not is_contaminated(text):
            set_paragraph_text(p, re.sub(r"\bconcept\b", "theme", text, flags=re.IGNORECASE))


def validate_document(doc: Document) -> list[str]:
    issues = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text
        for term in VALIDATION_TERMS:
            if term.lower() in t.lower():
                issues.append(f"Para {i}: found '{term}' -> {t[:80]}")
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for term in VALIDATION_TERMS:
                    if term.lower() in cell.text.lower():
                        issues.append(f"Table {ti} R{ri}C{ci}: found '{term}'")
    return issues


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(f"Source report not found: {SOURCE}")

    doc = Document(str(SOURCE))
    paragraphs = doc.paragraphs

    # Clear contaminated paragraphs
    contaminated = 0
    for p in paragraphs:
        if is_contaminated(p.text):
            clear_paragraph(p)
            contaminated += 1

    clean_tables(doc)

    # Replace key sections
    replace_section(paragraphs, "1.3 Objectives", ["1.4 Scope"], OBJECTIVES_TEXT)
    replace_section(paragraphs, "1.6 Proposed System", ["1.7"], PROPOSED_SYSTEM_TEXT)
    replace_section(paragraphs, "3.2 Non-Functional Requirements", ["3.3 Hardware"], NFR_TEXT)
    replace_section(
        paragraphs,
        "4.2 Methodology and Mathematical Formulation",
        ["Chapter 5"],
        MATH_TEXT,
    )

    fix_literature_concept_wording(paragraphs)

    for p in paragraphs:
        if "core concepts" in p.text.lower():
            set_paragraph_text(p, re.sub(r"core concepts", "core themes", p.text, flags=re.IGNORECASE))

    removed_dupes = remove_duplicates(paragraphs)
    removed_empty = remove_empty_paragraphs(doc)
    paragraphs = doc.paragraphs

    # Format all non-empty paragraphs
    for p in paragraphs:
        if p.text.strip():
            format_paragraph(p)

    # Second validation pass — clear any remaining hits
    issues = validate_document(doc)
    for p in paragraphs:
        for term in VALIDATION_TERMS:
            if term.lower() in p.text.lower():
                # Rewrite "Concept" only in literature theme lines already fixed; clear other hits
                if term == "Concept" and "Core Theme:" in p.text:
                    continue
                if is_contaminated(p.text) or term in ("ByteDaily", "Mastery", "SM-2", "Whisper"):
                    clear_paragraph(p)

    issues = validate_document(doc)
    if issues:
        print("VALIDATION WARNINGS (review manually):")
        for issue in issues[:30]:
            print(" ", issue)
    else:
        print("Validation passed: zero forbidden terms.")

    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")
    print(f"Contaminated paragraphs cleared: {contaminated}")
    print(f"Duplicate paragraphs removed: {removed_dupes}")
    print(f"Empty paragraphs deleted: {removed_empty}")


if __name__ == "__main__":
    main()
